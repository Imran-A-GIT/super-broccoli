from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..models import ContactInfo


def _set_font(run, size: int = 11, bold: bool = False, color: tuple | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size=11, bold=True, color=(31, 73, 125))
    # Add bottom border via paragraph XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_contact_block(doc: Document, name: str, contact_line: str, linkedin: str | None):
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    _set_font(run, size=16, bold=True)
    p.paragraph_format.space_after = Pt(2)

    # Contact line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(contact_line)
    _set_font(run2, size=10)
    p2.paragraph_format.space_after = Pt(2)

    if linkedin:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(linkedin)
        _set_font(run3, size=10)
        p3.paragraph_format.space_after = Pt(4)


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_font(run, size=10)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def _add_body_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=10, bold=bold)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def generate_docx(optimized_text: str, contact: ContactInfo) -> bytes:
    doc = Document()

    # Remove default margins for a cleaner look
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # Split on ## section markers
    raw_sections = re.split(r"\n## ", "\n" + optimized_text)
    raw_sections = [s for s in raw_sections if s.strip()]

    name = contact.name or "Candidate"
    contact_parts = [contact.email, contact.phone, contact.location]
    contact_line = " | ".join(p for p in contact_parts if p)

    first_section_done = False

    for raw in raw_sections:
        lines = raw.strip().splitlines()
        if not lines:
            continue

        header = lines[0].strip().upper()
        body_lines = lines[1:]

        if "CONTACT" in header:
            # Use our structured contact data for better formatting
            _add_contact_block(doc, name, contact_line, contact.linkedin)
            first_section_done = True
            continue

        if not first_section_done:
            # If no CONTACT section found, add contact block first
            _add_contact_block(doc, name, contact_line, contact.linkedin)
            first_section_done = True

        _add_heading(doc, header)

        for line in body_lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("### "):
                # Job title line
                _add_body_para(doc, stripped[4:], bold=True)
            elif re.match(r"^[-•*]\s+", stripped):
                # Bullet point
                bullet_text = re.sub(r"^[-•*]\s+", "", stripped)
                _add_bullet(doc, bullet_text)
            elif stripped.startswith("**") and stripped.endswith("**"):
                _add_body_para(doc, stripped.strip("*"), bold=True)
            else:
                _add_body_para(doc, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
