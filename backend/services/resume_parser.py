"""
Resume parser — handles PDF, DOCX, and plain text.

PDF: pdfplumber primary extraction + word-object fallback for text boxes / complex layouts.
DOCX: paragraphs + tables (two-column layouts) + XML text-box extraction + headers/footers.
Text: comprehensive regex-based section detection and structured data extraction.
"""
from __future__ import annotations

import io
import re
from collections import defaultdict
from typing import Optional

import pdfplumber
from docx import Document

from ..models import ContactInfo, Education, ResumeData, WorkExperience

# ─────────────────────────── SECTION HEADERS ────────────────────────────────

_SECTION_KEYS: dict[str, str] = {
    "SUMMARY": (
        r"SUMMARY|OBJECTIVE|PROFESSIONAL\s+SUMMARY|CAREER\s+OBJECTIVE|PROFILE|"
        r"PROFESSIONAL\s+PROFILE|EXECUTIVE\s+SUMMARY|CAREER\s+PROFILE|ABOUT\s+ME|"
        r"OVERVIEW|CAREER\s+OVERVIEW|PERSONAL\s+STATEMENT"
    ),
    "EXPERIENCE": (
        r"(?:WORK\s+|PROFESSIONAL\s+|RELATED\s+|RELEVANT\s+|CAREER\s+)?EXPERIENCE|"
        r"EMPLOYMENT(?:\s+HISTORY)?|WORK\s+HISTORY|CAREER\s+HISTORY|POSITIONS?\s+HELD|"
        r"PROFESSIONAL\s+BACKGROUND|WORK\s+BACKGROUND"
    ),
    "VOLUNTEER": (
        r"VOLUNTEER(?:ING|EER\s+EXPERIENCE)?(?:\s+EXPERIENCE)?|"
        r"COMMUNITY\s+(?:SERVICE|INVOLVEMENT|EXPERIENCE)|"
        r"CIVIC\s+(?:EXPERIENCE|INVOLVEMENT)|NON[\-\s]?PROFIT(?:\s+EXPERIENCE)?"
    ),
    "SKILLS": (
        r"(?:TECHNICAL\s+|CORE\s+|KEY\s+|PROFESSIONAL\s+|ADDITIONAL\s+)?SKILLS?|"
        r"COMPETENC(?:Y|IES)|AREAS?\s+OF\s+EXPERTISE|EXPERTISE|"
        r"TOOLS?(?:\s*(?:&|AND)\s*TECHNOLOGIES?)?|TECHNICAL\s+PROFICIENCIES?|"
        r"TECHNICAL\s+ABILITIES|PROFICIENCIES?|TECHNOLOGIES|QUALIFICATIONS?"
    ),
    "EDUCATION": (
        r"EDUCATION(?:AL\s+(?:BACKGROUND|QUALIFICATIONS?))?|"
        r"ACADEMIC\s+(?:BACKGROUND|QUALIFICATIONS?|EXPERIENCE)|"
        r"DEGREES?|SCHOOLING|ACADEMIC\s+TRAINING"
    ),
    "CERTIFICATIONS": (
        r"CERTIFICATIONS?|LICENS(?:E|ES|ING)|CREDENTIALS?|"
        r"PROFESSIONAL\s+CERTIFICATIONS?|LICENSES?\s*(?:&|AND)\s*CERTIFICATIONS?|"
        r"CONTINUING\s+EDUCATION|PROFESSIONAL\s+DEVELOPMENT|TRAINING"
    ),
    "PROJECTS": (
        r"(?:KEY\s+|PERSONAL\s+|PROFESSIONAL\s+|SELECTED\s+|NOTABLE\s+|RELEVANT\s+)?PROJECTS?(?:\s+EXPERIENCE)?|"
        r"PORTFOLIO|NOTABLE\s+WORK|PROJECT\s+HIGHLIGHTS?"
    ),
    "INTERESTS": (
        r"INTERESTS?|HOBBIES?(?:\s*(?:&|AND)\s*INTERESTS?)?|"
        r"PERSONAL\s+INTERESTS?|ACTIVITIES|EXTRACURRICULARS?"
    ),
    "LANGUAGES": r"LANGUAGES?(?:\s+(?:SKILLS?|PROFICIENCY|ABILITIES))?",
    "AWARDS": (
        r"AWARDS?(?:\s*(?:&|AND)\s*(?:RECOGNITION|HONORS?))?|ACHIEVEMENTS?|"
        r"ACCOMPLISHMENTS?|HONORS?|RECOGNITION|DISTINCTIONS?|ACCOLADES?"
    ),
    "PUBLICATIONS": r"PUBLICATIONS?|RESEARCH(?:\s+(?:AND|&)\s+PUBLICATIONS?)?|PAPERS?|PRESENTATIONS?|PATENTS?",
    "REFERENCES": r"REFERENCES?",
}

_SECTION_RE = re.compile(
    r"^(?:" + "|".join(f"(?:{p})" for p in _SECTION_KEYS.values()) + r")[\s:]*$",
    re.IGNORECASE,
)


def _canonical_section(text: str) -> str:
    """Map a section header line to its canonical key."""
    t = text.strip().rstrip(":").upper()
    for key, pattern in _SECTION_KEYS.items():
        if re.fullmatch(pattern, t, re.IGNORECASE):
            return key
    return t


# ─────────────────────────── CONTACT PATTERNS ───────────────────────────────

_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"(\+?1[\s\-.]?)?\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}")
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
_URL_RE = re.compile(
    r"https?://\S+|www\.\S+|\S+\.(?:com|org|net|io|co|edu|gov)/\S*",
    re.IGNORECASE,
)
_CORP_SUFFIX_RE = re.compile(
    r"\b(Inc\.?|Corp\.?|LLC|Ltd\.?|Corporation|Company|Technologies?|Systems?|"
    r"International|Group|Associates?|Services?|Solutions?|Holdings?|Enterprises?|"
    r"Consulting|Industries|Partners?|Global)\b",
    re.IGNORECASE,
)

# ─────────────────────────── DATE PATTERNS ──────────────────────────────────

_MONTH = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
)
_YEAR = r"\d{4}"
_ABBR_YEAR = r"'\d{2}"  # abbreviated year with apostrophe: '21, '19 (straight or smart)
_PRESENT = r"(?:Present|Current|Now|Today|Ongoing|To\s+Date)"

# Date token covers: "Aug 2021", "Aug '21", "August, 2021", "2021",
# "08/2021", "2021-08", "Q1 2022" (quarterly), standalone 4-digit year
_DATE_TOKEN = rf"(?:{_MONTH}\s*[,.]?\s*(?:{_YEAR}|{_ABBR_YEAR})|Q[1-4]\s*{_YEAR}|{_YEAR}|\d{{1,2}}[/\-]{_YEAR})"

_DATE_RANGE_RE = re.compile(
    rf"(?P<start>{_DATE_TOKEN})\s*[-–—to/]+\s*(?P<end>{_DATE_TOKEN}|{_PRESENT})",
    re.IGNORECASE,
)

_DEGREE_RE = re.compile(
    r"\b(Bachelor(?:'?s)?|Master(?:'?s)?|MBA|M\.B\.A\.?|PhD|Ph\.D\.?|"
    r"Associate(?:'?s)?|B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|B\.?Sc\.?|M\.?Sc\.?|"
    r"B\.?Eng\.?|M\.?Eng\.?|Doctor(?:ate|al)?|Diploma|GED|"
    r"A\.?A\.?|A\.?S\.?|D\.B\.A\.?|J\.D\.?|LL\.?M\.?)\b",
    re.IGNORECASE,
)

_CERT_RE = re.compile(
    r"\b(Certified|Certification|Certificate|CompTIA|CISSP|CISM|CISA|CEH|CHFI|"
    r"AWS\s+(?:Certified|Solutions)|Azure\s+(?:Certified|Administrator|Developer)|"
    r"GCP|Google\s+Cloud|PMP|ITIL|PRINCE2|Scrum|PMI|CAPM|CSM|CSD|"
    r"Security\+|Network\+|A\+|CySA\+|CCNA|CCNP|CCIE|MCSA|MCSE|MCTS|MCITP|"
    r"SHRM|CPA|CFA|CIPP|OSCP|GIAC|GPEN|GSEC|Six\s+Sigma|Lean|PCI(?:-DSS)?|"
    r"TOGAF|COBIT|ISO\s+27001|SOC\s*2|HIPAA|FINRA|Series\s+\d+)\b",
    re.IGNORECASE,
)

_JOB_TITLE_RE = re.compile(
    r"\b(Analyst|Developer|Manager|Designer|Director|Coordinator|Specialist|"
    r"Consultant|Advisor|Administrator|Architect|Lead|Senior|Junior|Executive|"
    r"Officer|Supervisor|Technician|Programmer|Scientist|Representative|"
    r"Inspector|Auditor|Researcher|Strategist|President|Intern|co-op|"
    r"Assistant|Processor|Contractor)\b",
    re.IGNORECASE,
)

_BULLET_RE = re.compile(
    # Common bullet/symbol characters used across resume templates
    r"^[•‣·▪●➢✓►▸❖▪▸◆►○●•·→⟩›»⦿⁃∙◉◎✦✧✗✘✔✕✖"
    r"\-\*]\s+(.+)"
)
_NUMBERED_BULLET_RE = re.compile(r"^\d+[.)]\s+(.+)")
# Lines that are purely decorative dividers (─, ═, ─, ━, =, _, -)
_DIVIDER_RE = re.compile(r"^[\-─═━_=\s]{4,}$")


# ─────────────────────────── PDF EXTRACTION ─────────────────────────────────

def _find_column_split(words: list[dict], page_width: float) -> Optional[float]:
    """Find the x-coordinate gap between a left sidebar and right main column.

    Only fires when there is a clear whitespace channel between two distinct
    text columns (gap >= 5% of page width and left column width <= 42% of page).
    This conservative threshold avoids false-positives on single-column resumes
    that have indented bullets or centred headers.
    """
    lo = page_width * 0.10
    hi = page_width * 0.48

    # Collect distinct word-start x-positions within the search band
    xs = sorted({
        int(float(w.get("x0", 0)))
        for w in words
        if lo <= float(w.get("x0", 0)) <= hi
    })
    if len(xs) < 4:  # need enough diversity to distinguish two populations
        return None

    max_gap = 0
    split_x: Optional[float] = None
    for i in range(1, len(xs)):
        gap = xs[i] - xs[i - 1]
        if gap > max_gap:
            max_gap = gap
            split_x = (xs[i] + xs[i - 1]) / 2.0

    # Gap must be >= 5% of page width (≈ 30pt on US letter) to be a real column gap.
    # Single-column indentation typically creates gaps of 10-20pt which are filtered out.
    if split_x is None or max_gap < page_width * 0.05:
        return None

    # Sanity: left "column" must not extend past 45% of page width
    left_max_x1 = max(
        (float(w.get("x1", w.get("x0", 0))) for w in words if float(w.get("x0", 0)) < split_x),
        default=0,
    )
    if left_max_x1 > page_width * 0.45:
        return None

    # Both columns need enough words to be real content blocks
    left_count = sum(1 for w in words if float(w.get("x0", 0)) < split_x)
    right_count = sum(1 for w in words if float(w.get("x0", 0)) >= split_x)
    if left_count < 8 or right_count < 20:
        return None

    return split_x


def _extract_page_text(page) -> str:
    """Extract text from one PDF page, handling two-column sidebar layouts.

    Uses word x-coordinates (not page crops) to separate columns, so words
    that straddle the typical 35% crop boundary are never cut in half.
    Right column (main content / work experience) is emitted first so that
    section headers appear before sidebar items in the assembled text.
    """
    try:
        w = float(page.width)
        words = page.extract_words(x_tolerance=3, y_tolerance=3)
        if words and len(words) > 15:
            split_x = _find_column_split(words, w)
            if split_x is not None:
                left_words = [wd for wd in words if float(wd.get("x0", 0)) < split_x]
                right_words = [wd for wd in words if float(wd.get("x0", 0)) >= split_x]
                left_text = _words_to_lines(left_words).strip()
                right_text = _words_to_lines(right_words).strip()
                if len(left_text) > 80 and len(right_text) > 100:
                    # Determine which column is the sidebar (shorter, contains
                    # sidebar-typical sections: INTERESTS/SKILLS/LANGUAGES).
                    # Emit the main-content column first so EXPERIENCE headers
                    # appear before sidebar section headers.
                    _SIDEBAR_SECTIONS = re.compile(
                        r"\b(INTERESTS?|HOBBIES?|SKILLS?|LANGUAGES?|CONTACT|SOCIAL)\b",
                        re.IGNORECASE,
                    )
                    left_is_sidebar = bool(_SIDEBAR_SECTIONS.search(left_text)) or (
                        len(left_text) < len(right_text) * 0.6
                    )
                    if left_is_sidebar:
                        return right_text + "\n\n" + left_text
                    else:
                        return left_text + "\n\n" + right_text
    except Exception:
        pass

    # Standard single-column extraction
    text = page.extract_text(x_tolerance=3, y_tolerance=3)
    if text and len(text.strip()) > 30:
        return text

    # Word-object fallback for complex layouts
    words = page.extract_words(
        x_tolerance=5, y_tolerance=5,
        keep_blank_chars=False,
        use_text_flow=True,
    )
    return _words_to_lines(words) if words else ""


async def parse_pdf(file_bytes: bytes) -> ResumeData:
    """Extract text from PDF and parse into structured resume data."""
    texts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = _extract_page_text(page)
            if page_text:
                texts.append(page_text)

    full_text = "\n\n".join(texts)
    if not full_text.strip():
        raise ValueError(
            "Could not extract text from PDF. "
            "Ensure it is a text-based PDF (not a scanned image) or paste your resume directly."
        )
    return _parse_text(full_text)


def _words_to_lines(words: list[dict]) -> str:
    """Reconstruct text lines from positioned word objects, handling multi-column layouts."""
    if not words:
        return ""
    # Group by y-position (3px tolerance to merge words on the same line)
    by_line: dict[int, list[dict]] = defaultdict(list)
    for w in words:
        y_key = round(float(w.get("top", 0)) / 3) * 3
        by_line[y_key].append(w)

    lines: list[str] = []
    for y in sorted(by_line):
        row = sorted(by_line[y], key=lambda w: float(w.get("x0", 0)))
        lines.append(" ".join(w["text"] for w in row))
    return "\n".join(lines)


# ─────────────────────────── DOCX EXTRACTION ────────────────────────────────

async def parse_docx(file_bytes: bytes) -> ResumeData:
    """Extract text from DOCX including text boxes, tables, and headers/footers."""
    doc = Document(io.BytesIO(file_bytes))
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    all_parts: list[str] = []

    # 1. Headers/footers first (often contain name + contact info)
    hf_lines: list[str] = []
    try:
        for section in doc.sections:
            for hf in (section.header, section.footer):
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if t:
                        hf_lines.append(t)
    except Exception:
        pass
    if hf_lines:
        all_parts.append("\n".join(hf_lines))

    # 2. Regular paragraphs (main body text)
    para_lines: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            para_lines.append(t)
    if para_lines:
        all_parts.append("\n".join(para_lines))

    # 3. Tables — many resumes use tables for two-column layouts
    seen_cells: set[str] = set()
    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_cells: list[str] = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct and ct not in seen_cells:
                    seen_cells.add(ct)
                    row_cells.append(ct)
            if row_cells:
                # Join cells with separator; each cell may be a separate column
                table_lines.append("  |  ".join(row_cells))
    if table_lines:
        all_parts.append("\n".join(table_lines))

    # 4. XML text boxes (<w:txbxContent>) — floating text frames
    txbx_lines: list[str] = []
    try:
        for txbx in doc.element.body.iter(f"{{{W}}}txbxContent"):
            for p in txbx.iter(f"{{{W}}}p"):
                text = "".join(t.text or "" for t in p.iter(f"{{{W}}}t")).strip()
                if text:
                    txbx_lines.append(text)
    except Exception:
        pass
    if txbx_lines:
        all_parts.append("\n".join(txbx_lines))

    # 5. Drawing / SmartArt inline text
    try:
        drawing_texts: list[str] = []
        for elem in doc.element.body.iter():
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "t" and elem.text and elem.text.strip():
                # Only capture if not already in regular paragraphs
                if elem.text.strip() not in "\n".join(all_parts):
                    drawing_texts.append(elem.text.strip())
        if drawing_texts:
            all_parts.append("\n".join(drawing_texts))
    except Exception:
        pass

    full_text = "\n\n".join(all_parts)
    return _parse_text(full_text)


# ─────────────────────────── PLAIN TEXT ─────────────────────────────────────

async def parse_text(raw_text: str) -> ResumeData:
    return _parse_text(raw_text)


# ─────────────────────────── MAIN PARSER ────────────────────────────────────

def _parse_text(text: str) -> ResumeData:
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Normalize Unicode punctuation: smart quotes/apostrophes → straight ASCII.
    # PDFs often embed curly apostrophes (U+2018/2019) which break date regex like "Aug '21".
    text = (
        text
        .replace("‘", "'").replace("’", "'")  # left/right single quotes
        .replace("“", '"').replace("”", '"')  # left/right double quotes
        .replace("ʼ", "'").replace("′", "'")  # modifier apostrophe, prime
        .replace("–", "–").replace("—", "—")  # en-dash, em-dash (keep as-is but ensure)
    )
    lines = [ln.rstrip() for ln in text.splitlines()]

    contact = _extract_contact(lines, text)
    sections = _split_sections(lines)

    return ResumeData(
        contact=contact,
        summary=_extract_summary(sections),
        experience=_extract_jobs(sections, "EXPERIENCE"),
        volunteer_experience=_extract_jobs(sections, "VOLUNTEER"),
        skills=_extract_skills(sections),
        education=_extract_education(sections),
        certifications=_extract_certifications(sections, text),
        projects=_extract_list_section(sections, "PROJECTS"),
        interests=_extract_list_section(sections, "INTERESTS"),
        languages=_extract_list_section(sections, "LANGUAGES"),
        awards=_extract_list_section(sections, "AWARDS"),
        raw_text=text,
    )


# ─────────────────────────── CONTACT ────────────────────────────────────────

def _extract_contact(lines: list[str], full_text: str) -> ContactInfo:
    email = (m.group() if (m := _EMAIL_RE.search(full_text)) else None)
    phone = (m.group() if (m := _PHONE_RE.search(full_text)) else None)
    linkedin = (m.group() if (m := _LINKEDIN_RE.search(full_text)) else None)

    # Name detection — three passes:
    # Pass 1: "Name | email | phone" combined header — name is ALWAYS the first token
    name: Optional[str] = None
    for line in lines[:3]:
        stripped = line.strip()
        if re.search(r"\s[|·•]\s", stripped) or stripped.count("  ") >= 1:
            first_token = re.split(r"\s*[|·•]\s*|\s{2,}", stripped)[0].strip()
            if (
                first_token
                and not _EMAIL_RE.search(first_token)
                and not _PHONE_RE.search(first_token)
                and not _URL_RE.search(first_token)
                and 2 < len(first_token) < 50
                and not re.match(r"^\d", first_token)
                and not (first_token.isupper() and (len(first_token.split()) > 2 or _CORP_SUFFIX_RE.search(first_token)))
            ):
                name = first_token
                break

    # Pass 2 (fallback): first short standalone non-contact, non-header line in first 12 lines
    if not name:
        for line in lines[:12]:
            stripped = line.strip()
            if (
                stripped
                and not _EMAIL_RE.search(stripped)
                and not _PHONE_RE.search(stripped)
                and not _URL_RE.search(stripped)
                and not _SECTION_RE.match(stripped)
                and not stripped.startswith("+")
                and not re.match(r"^\d", stripped)
                and "|" not in stripped
                and not _DATE_RANGE_RE.search(stripped)
                and 2 < len(stripped) < 50
                and not (stripped.isupper() and (len(stripped.split()) > 2 or _CORP_SUFFIX_RE.search(stripped)))
            ):
                name = stripped
                break

    # Pass 3 (extended fallback): scan all lines for a name-like pattern.
    # Handles two-column PDFs where the name appears far into the extracted text
    # (e.g. left-column name gets read after right-column work experience).
    if not name:
        _PREPOSITIONS = frozenset({"of", "and", "or", "at", "in", "the", "for", "by", "to", "a", "an"})
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if (_EMAIL_RE.search(stripped) or _PHONE_RE.search(stripped)
                    or _URL_RE.search(stripped) or stripped.startswith("+")
                    or re.match(r"^\d", stripped) or "|" in stripped
                    or _BULLET_RE.match(stripped) or _DATE_RANGE_RE.search(stripped)
                    or _SECTION_RE.match(stripped.rstrip(":"))
                    or _CORP_SUFFIX_RE.search(stripped)
                    or len(stripped) > 45):
                continue
            # Split and strip trailing credentials (short all-caps like CSM, PMP, MBA)
            raw_tokens = re.split(r"[\s,]+", stripped)
            name_tokens = [
                t for t in raw_tokens
                if t and not (len(re.sub(r"[®©™]", "", t)) <= 4 and re.sub(r"[®©™]", "", t).isupper())
            ]
            if 1 < len(name_tokens) <= 4:
                lower_toks = [t.lower().strip(".,®©™") for t in name_tokens]
                if not any(t in _PREPOSITIONS for t in lower_toks):
                    if all(re.match(r"^[A-Za-z\'\-]+$", t) for t in name_tokens):
                        name = " ".join(t.title() if t.isupper() else t for t in name_tokens)
                        break

    # Location: City, ST or City, Country
    location: Optional[str] = None
    for pattern in [
        re.compile(r"\b[A-Z][a-zA-Z\s\-]+,\s*[A-Z]{2}\b"),
        re.compile(r"\b[A-Z][a-zA-Z\s\-]+,\s*[A-Z][a-zA-Z\s]+\b"),
    ]:
        if m := pattern.search(full_text):
            candidate = m.group().strip()
            if len(candidate) < 50 and not _EMAIL_RE.search(candidate):
                location = candidate
                break

    return ContactInfo(
        name=name, email=email, phone=phone, linkedin=linkedin, location=location
    )


# ─────────────────────────── SECTION SPLITTING ──────────────────────────────

def _split_sections(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {"HEADER": []}
    current = "HEADER"
    for line in lines:
        stripped = line.strip().rstrip(":")

        # Skip purely decorative divider lines (═══, ───, ___, ===, etc.)
        if _DIVIDER_RE.match(stripped):
            continue

        # Strip leading/trailing decorative symbols (some templates wrap headers
        # in symbols: "── EXPERIENCE ──" or "● SKILLS ●" or ">>> EDUCATION <<<")
        normalized = re.sub(
            r"^[\-─═━_=•·▪▸►●◆◉○→»«<>\s]+|[\-─═━_=•·▪▸►●◆◉○→»«<>\s]+$",
            "", stripped,
        ).strip()
        # Strip continuation markers: "WORK EXPERIENCE (cont.)" → "WORK EXPERIENCE"
        normalized = re.sub(
            r"\s*\(cont(?:inued)?\.?\)\s*$|\s*\(page\s*\d+\)\s*$|\s*-\s*cont(?:inued)?\.?\s*$",
            "", normalized, flags=re.IGNORECASE,
        ).strip()

        if normalized and (_SECTION_RE.match(normalized + " ") or _SECTION_RE.match(normalized)):
            key = _canonical_section(normalized)
            current = key
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return sections


# ─────────────────────────── SUMMARY ────────────────────────────────────────

def _extract_summary(sections: dict[str, list[str]]) -> Optional[str]:
    for key in sections:
        if key == "SUMMARY":
            text = "\n".join(ln for ln in sections[key] if ln.strip())
            return text.strip() or None
    return None


# ─────────────────────────── EXPERIENCE ─────────────────────────────────────

def _extract_jobs(sections: dict[str, list[str]], section_type: str) -> list[WorkExperience]:
    """Collect and parse all lines from sections matching the given type."""
    lines: list[str] = []
    for key in sections:
        if key == section_type:
            lines.extend(sections[key])

    if section_type == "EXPERIENCE":
        # Two-column PDF fix: job entries that appear before the first section
        # header land in HEADER.  Include them if they contain date ranges.
        header_lines = sections.get("HEADER", [])
        if any(_DATE_RANGE_RE.search(l) for l in header_lines):
            lines = header_lines + lines

        # Safety net: if co-op / older roles landed in INTERESTS or PROJECTS
        # (due to page-boundary section attribution), absorb them too.
        for extra_key in ("INTERESTS", "PROJECTS", "AWARDS"):
            extra = sections.get(extra_key, [])
            if any(_DATE_RANGE_RE.search(l) for l in extra):
                lines = lines + extra

    if not lines:
        return []

    jobs = _parse_job_blocks(lines)

    # Deduplicate by (title[:30], company[:20], start_date) — guards against
    # the same entry being picked up from both HEADER and EXPERIENCE.
    seen: set[tuple[str, str, str]] = set()
    unique: list[WorkExperience] = []
    for j in jobs:
        key = (j.title.lower()[:30], j.company.lower()[:20], (j.start_date or "").lower())
        if key not in seen:
            seen.add(key)
            unique.append(j)
    return unique


def _parse_job_blocks(lines: list[str]) -> list[WorkExperience]:
    """
    Parse a flat list of text lines into WorkExperience objects.

    Strategy: anchor on date-range lines, look backwards for title/company,
    look forwards for bullet points.
    """
    # Locate every line that contains a date range
    date_hits: list[tuple[int, re.Match]] = []
    for i, line in enumerate(lines):
        m = _DATE_RANGE_RE.search(line.strip())
        if m:
            date_hits.append((i, m))

    if not date_hits:
        return _parse_jobs_undated(lines)

    jobs: list[WorkExperience] = []

    for job_idx, (date_line_i, date_match) in enumerate(date_hits):
        start_date = date_match.group("start").strip()
        end_date = date_match.group("end").strip()

        # ── Determine header region for this job ──────────────────────────
        # Starts right after the previous job's date line (or beginning of list)
        prev_date_line = date_hits[job_idx - 1][0] if job_idx > 0 else -1
        header_lines: list[str] = []

        for k in range(prev_date_line + 1, date_line_i):
            stripped = lines[k].strip()
            if not stripped:
                continue
            # Skip lines that are bullet points (they belong to the previous job)
            if _BULLET_RE.match(stripped) or _NUMBERED_BULLET_RE.match(stripped):
                continue
            # Skip lines that are section headers
            if _SECTION_RE.match(stripped.rstrip(":")):
                break
            # Skip lines that ARE another date range (shouldn't happen, but guard)
            if _DATE_RANGE_RE.search(stripped):
                break
            header_lines.append(stripped)

        # Also capture any title/company info embedded on the date line itself
        remainder_raw = _DATE_RANGE_RE.sub("", lines[date_line_i]).strip()
        # Strip dashes but preserve leading | (needed to detect continuation/promotion entries)
        remainder = remainder_raw.lstrip("–—").strip()
        if remainder:
            header_lines.append(remainder)

        # Detect "continuation" entries: multiple roles at the same company (promotions).
        # Indicated by the date line or remainder starting with a pipe character.
        is_continuation = (
            lines[date_line_i].strip().startswith("|")
            or remainder_raw.startswith("|")
            or (header_lines and header_lines[0].lstrip().startswith("|"))
        )

        title, company = _parse_title_company(header_lines)

        # For promotion/continuation entries inherit the company from the previous role
        if is_continuation and (not company or company == "Company") and jobs:
            company = jobs[-1].company

        # ── Collect bullet points ─────────────────────────────────────────
        next_date_line = (
            date_hits[job_idx + 1][0] if job_idx + 1 < len(date_hits) else len(lines)
        )
        bullets: list[str] = []
        for k in range(date_line_i + 1, next_date_line):
            stripped = lines[k].strip()
            if not stripped:
                continue
            b = _BULLET_RE.match(stripped) or _NUMBERED_BULLET_RE.match(stripped)
            if b:
                bullets.append(b.group(1).strip())
            elif (
                len(stripped) > 20
                and not _DATE_RANGE_RE.search(stripped)
                and not stripped.isupper()
                and not _SECTION_RE.match(stripped.rstrip(":"))
                and k > date_line_i + 2  # skip location line immediately after date
            ):
                # Long prose line — treat as a non-bulleted description
                bullets.append(stripped)

        jobs.append(
            WorkExperience(
                title=title or "Position",
                company=company or "Company",
                start_date=start_date,
                end_date=end_date,
                bullets=bullets,
            )
        )

    return _sort_by_recency(jobs)


def _parse_jobs_undated(lines: list[str]) -> list[WorkExperience]:
    """Fallback for experience sections that have no date ranges at all."""
    jobs: list[WorkExperience] = []
    title = company = ""
    bullets: list[str] = []

    def _flush() -> None:
        if title or company:
            jobs.append(
                WorkExperience(
                    title=title.strip() or "Position",
                    company=company.strip() or "Company",
                    bullets=list(bullets),
                )
            )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        b = _BULLET_RE.match(stripped) or _NUMBERED_BULLET_RE.match(stripped)
        if b:
            bullets.append(b.group(1).strip())
        elif not title:
            _flush()
            title, company = stripped, ""
            bullets = []
        elif not company and len(stripped) < 80:
            company = stripped
        else:
            bullets.append(stripped)

    _flush()
    return jobs


def _parse_title_company(header_lines: list[str]) -> tuple[str, str]:
    """Determine job title and company name from a list of header lines."""
    title = company = ""
    if not header_lines:
        return title, company

    # Strip noise: dividers, section-header-like lines, and very short fragments
    # (e.g. "(cont.)" lines that can appear from page-continuation markers)
    _NOISE_RE = re.compile(
        r"^\s*(\(cont(?:inued)?\.?\)|cont\.?|page\s*\d+|"
        r"[\-─═━_=]{2,})\s*$", re.IGNORECASE
    )
    clean = [l for l in header_lines if l.strip() and not _NOISE_RE.match(l.strip())]
    if not clean:
        return title, company

    # Check for combined separators on a single line
    for line in clean:
        # "Title at Company"
        if re.search(r"\bat\b", line, re.IGNORECASE) and not _DATE_RANGE_RE.search(line):
            parts = re.split(r"\s+at\s+", line, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                return parts[0].strip(), parts[1].strip()
        # "Title @ Company"
        if " @ " in line:
            parts = line.split(" @ ", 1)
            return parts[0].strip(), parts[1].strip()
        # Pipe-separated: "Company | Title" / "Title | Company" / "Title | Company | Location"
        if "|" in line and not _DATE_RANGE_RE.search(line):
            if line.lstrip().startswith("|"):
                # Leading pipe → continuation entry (same company, different role)
                return line.lstrip().lstrip("|").strip(), ""
            segments = [s.strip() for s in re.split(r"\s*\|\s*", line) if s.strip()]
            if len(segments) >= 2:
                left, right = segments[0], segments[1]
                right_alpha = re.sub(r"[^A-Za-z]", "", right)
                # Right side ALL CAPS → it's the title (resume format: Company | TITLE)
                if right_alpha and right_alpha.isupper():
                    return right, left
                # Right has title keywords, left doesn't → Company | Title
                if _JOB_TITLE_RE.search(right) and not _JOB_TITLE_RE.search(left):
                    return right, left
                # Left has corporate suffix → left is company
                if _CORP_SUFFIX_RE.search(left):
                    return right, left
                return left, right
        # "Title — Company" or "Title – Company" (em/en dash)
        for sep in (" — ", " – ", " ─ "):
            if sep in line and not _DATE_RANGE_RE.search(line):
                parts = line.split(sep, 1)
                if parts[0].strip() and parts[1].strip():
                    return parts[0].strip(), parts[1].strip()
        # "Title - Company" (hyphen, only if both sides have 3+ chars)
        if " - " in line and not _DATE_RANGE_RE.search(line):
            parts = line.split(" - ", 1)
            if len(parts[0].strip()) > 2 and len(parts[1].strip()) > 2:
                return parts[0].strip(), parts[1].strip()
        # "Title, Company" — require substantial second part to avoid splitting
        # "Name, Credential" or "City, State"
        if "," in line and not _DATE_RANGE_RE.search(line):
            parts = line.split(",", 1)
            second = parts[1].strip()
            if second and len(second) > 10 and not re.match(r"^\s*(Inc|LLC|Ltd|Corp|[A-Z]{2})\b", second, re.I):
                return parts[0].strip(), second

    # Multi-line header: two separate lines for title and company
    if len(clean) >= 2:
        first, second = clean[0], clean[1]
        first_upper = re.sub(r"[^A-Za-z]", "", first).isupper()
        second_upper = re.sub(r"[^A-Za-z]", "", second).isupper()
        # Format: Company (ALL CAPS or has corp suffix) on line 1, Title on line 2
        if (first_upper and len(first) > 3) or _CORP_SUFFIX_RE.search(first):
            company = first.title() if first_upper else first
            title = second
        # Format: Title on line 1, Company (has corp suffix or is ALL CAPS) on line 2
        elif (second_upper and len(second) > 3) or _CORP_SUFFIX_RE.search(second):
            title = first
            company = second.title() if second_upper else second
        else:
            # Default: first line = title, second non-URL/non-numeric line = company
            title = first
            for hl in clean[1:]:
                if not re.match(r"^\d", hl) and not hl.startswith("http") and len(hl) < 80:
                    company = hl
                    break
    elif clean:
        title = clean[0]

    return title.strip(), company.strip()


def _sort_by_recency(jobs: list[WorkExperience]) -> list[WorkExperience]:
    """Return jobs sorted most-recent first based on end date then start date."""
    def _year(s: Optional[str]) -> int:
        if not s:
            return 0
        if re.search(r"present|current|now|today|ongoing", s, re.IGNORECASE):
            return 9999
        m = re.search(r"\b(19|20)\d{2}\b", s)
        if m:
            return int(m.group())
        # Handle abbreviated years: '21 → 2021, '95 → 1995
        m = re.search(r"'(\d{2})\b", s)
        if m:
            yy = int(m.group(1))
            return 2000 + yy if yy <= 50 else 1900 + yy
        return 0

    return sorted(jobs, key=lambda j: (-_year(j.end_date), -_year(j.start_date)))


# ─────────────────────────── SKILLS ─────────────────────────────────────────

def _extract_skills(sections: dict[str, list[str]]) -> list[str]:
    for key in sections:
        if key == "SKILLS":
            raw = "\n".join(sections[key])
            # Handle "Category: skill1, skill2" patterns
            raw = re.sub(r"^[^\n:]{1,30}:\s*", "", raw, flags=re.MULTILINE)
            items = re.split(r"[,;|\n••\-\*▪]", raw)
            skills = [
                s.strip().strip(".")
                for s in items
                if s.strip() and len(s.strip()) > 1 and not _SECTION_RE.match(s.strip())
            ]
            return skills[:80]
    return []


# ─────────────────────────── EDUCATION ──────────────────────────────────────

def _extract_education(sections: dict[str, list[str]]) -> list[Education]:
    for key in sections:
        if key == "EDUCATION":
            edu_list: list[Education] = []
            lines = [ln.strip() for ln in sections[key] if ln.strip()]
            i = 0
            while i < len(lines):
                line = lines[i]
                if _DEGREE_RE.search(line):
                    degree = line
                    institution = ""
                    if i + 1 < len(lines):
                        institution = lines[i + 1]
                        i += 1
                    year_src = degree + " " + institution
                    ym = re.search(r"\b(19|20)\d{2}\b", year_src)
                    year = ym.group() if ym else None
                    edu_list.append(
                        Education(degree=degree, institution=institution, year=year)
                    )
                i += 1
            return edu_list
    return []


# ─────────────────────────── CERTIFICATIONS ─────────────────────────────────

def _extract_certifications(sections: dict[str, list[str]], full_text: str) -> list[str]:
    for key in sections:
        if key == "CERTIFICATIONS":
            lines = [ln.strip() for ln in sections[key] if ln.strip()]
            return [re.sub(r"^[••\-\*]\s*", "", ln) for ln in lines if ln]

    # Fallback: scan entire text for certification-like lines
    certs: list[str] = []
    for line in full_text.splitlines():
        stripped = line.strip()
        if _CERT_RE.search(stripped) and 5 < len(stripped) < 120:
            clean = re.sub(r"^[••\-\*]\s*", "", stripped)
            if clean and clean not in certs:
                certs.append(clean)
    return certs[:15]


# ─────────────────────────── GENERIC LIST SECTIONS ──────────────────────────

def _extract_list_section(sections: dict[str, list[str]], section_key: str) -> list[str]:
    """Extract a section as a flat list of strings (projects, interests, etc.)."""
    for key in sections:
        if key == section_key:
            raw = " ".join(ln.strip() for ln in sections[key] if ln.strip())
            # Split on common delimiters
            items = re.split(r"[,;|\n••\-\*▪]", raw)
            cleaned: list[str] = []
            for item in items:
                item = item.strip().strip(".")
                if item and 1 < len(item) < 120 and not _SECTION_RE.match(item):
                    cleaned.append(item)
            return cleaned
    return []
